import os
try:
    import docx
except ImportError:
    print("python-docx is not installed. Please install it first.")
    exit(1)

def main():
    doc = docx.Document()
    
    # Add title
    doc.add_heading("Lab Assignment - Source Code", 0)

    # We can add files manually or by pattern
    files_to_read = [
        "app.py",
        "models.py",
        "routes.py",
        "requirements.txt",
        "static/css/style.css",
        "static/js/main.js",
    ]
    
    # Add template files
    templates_dir = "templates"
    if os.path.exists(templates_dir):
        for file in sorted(os.listdir(templates_dir)):
            if file.endswith(".html"):
                files_to_read.append(os.path.join(templates_dir, file).replace('\\', '/'))
            
    for file_path in files_to_read:
        if os.path.exists(file_path):
            doc.add_heading(f"{file_path}", level=1)
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
                
                # Add content using a fixed-width style if possible, or just normal paragraph
                p = doc.add_paragraph(content)
                # Ensure spacing isn't too huge
                p.paragraph_format.space_after = docx.shared.Pt(0)
                
            except Exception as e:
                doc.add_paragraph(f"Error reading file: {e}")
                
    output_name = "Assignment_Code.docx"
    doc.save(output_name)
    print(f"Successfully created {output_name}")

if __name__ == "__main__":
    main()
