import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def main():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading("Campus Hobby Connector", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle = doc.add_paragraph("Comprehensive Deployment & Testing Guide")
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].bold = True
    
    doc.add_page_break()

    # Introduction
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph("This document serves as a comprehensive manual for deploying and testing the Campus Hobby Connector System. It outlines every necessary step to set up the environment, initialize the database, and run automated & manual tests for all system modules.")
    
    # Deployment Section
    doc.add_heading("2. Complete Deployment Steps", level=1)
    doc.add_paragraph("Follow these exact steps to deploy the application on a local development environment.")
    
    deploy_steps = [
        ("Step 2.1: Requirements Verification", 
         "Before beginning, verify that you have Python 3.8 or newer installed on your machine. You can check this by running `python --version` in your terminal. You will also need a terminal (Command Prompt, PowerShell, or Bash)."),
        ("Step 2.2: Virtual Environment Setup", 
         "To isolate the project's dependencies from your system's global Python packages, it is highly recommended to use a virtual environment. Inside the project folder, run:\n• Windows: `python -m venv .venv` followed by `.\\.venv\\Scripts\\activate`\n• macOS/Linux: `python3 -m venv .venv` followed by `source .venv/bin/activate`"),
        ("Step 2.3: Dependency Installation", 
         "The system relies on several external packages like Flask, SQLAlchemy, and Flask-SocketIO. Install them by running:\n`pip install -r requirements.txt`.\nWait for all installations to complete successfully without errors."),
        ("Step 2.4: Database Initialization", 
         "The application auto-configures a local SQLite database named `database.db`. When you run the application for the first time, models are automatically mapped to tables, and seed data (like default hobbies) is injected. To trigger this, no manual action is needed beyond starting the app."),
        ("Step 2.5: Starting the Application Server", 
         "Execute the command `python app.py`. You should see the terminal log: `🚀 Campus Hobby Connector is running!`. The application uses an event-loop server for SocketIO, binding typically to port 5000."),
        ("Step 2.6: Accessing the Portal", 
         "Open a web browser (Chrome or Firefox recommended) and navigate to `http://127.0.0.1:5000`. You will be greeted by the Landing Page.")
    ]
    
    for step_title, step_desc in deploy_steps:
        doc.add_heading(step_title, level=2)
        doc.add_paragraph(step_desc)

    doc.add_page_break()

    # Testing Section
    doc.add_heading("3. Elaborate Testing Guide", level=1)
    doc.add_paragraph("This section breaks down the functional testing for the four main modules: Authentication, Matchmaking, Real-time WebSockets Chat, and Administrative functions.")
    
    test_cases = [
        ("Test Module 1: User Authentication & Sessions", [
            "Objective: Verify users can securely register, login, and maintain sessions securely.",
            "Action 1 (Registration): Navigate to http://127.0.0.1:5000/register. Enter a username (e.g., 'Alice'), an email (e.g., 'alice@test.com'), and a password. Click Submit.",
            "Expected Result: The account is dynamically created in the database, the password is encrypted using Bcrypt, and you are redirected to the dashboard.",
            "Action 2 (Login verification): Log out from the navigation bar. Go to the Login page. Try to log in with an incorrect password.",
            "Expected Result: An 'Invalid Credentials' error message should be displayed.",
            "Action 3 (Valid Log in): Enter the correct email and password for Alice.",
            "Expected Result: You successfully authenticate and the session token is mapped."
        ]),
        ("Test Module 2: Hobby Configuration & State Storage", [
            "Objective: Ensure users can add default mapped hobbies and custom hobbies.",
            "Action 1: After logging in, navigate to the Profile & Hobbies section.",
            "Action 2: Update the User Bio by typing 'Hello, I love coding!' and saving.",
            "Action 3: From the Hobbies checklist, select 'Coding' and 'Photography'. In the Custom Hobbies input box, enter 'Skateboarding, Chess'. Press Save.",
            "Expected Result: The dashboard now displays the 4 hobbies perfectly rendered. Behind the scenes, 'Skateboarding' and 'Chess' were added to the Master Hobbies Table if they did not already exist, and mapped to the User via the UserHobby relationship table."
        ]),
        ("Test Module 3: Matchmaking Algorithms", [
            "Objective: Verify the Matchmaking Engine correctly identifies overlapping hobbies and returns a compatibility score.",
            "Action 1: Register a second account 'Bob' using an incognito or private window.",
            "Action 2: For Bob's hobbies, select 'Coding' and 'Chess'.",
            "Action 3: Navigate to the 'Matches' tab while logged in as Alice.",
            "Expected Result: Bob should appear in the match results with a score of approx 33-50% (since Alice has 4 hobbies and Bob has 2, matching exactly 2).",
            "Action 4: Verify the UI shows exactly which hobbies were common (Coding, Chess)."
        ]),
        ("Test Module 4: Real-Time Communication (Socket.IO)", [
            "Objective: Test full-duplex WebSocket connections for instantaneous messaging.",
            "Action 1: Keep Alice open in Window 1 and Bob in Window 2.",
            "Action 2: In Alice's window, go to Matches, and click 'Chat with Bob'.",
            "Action 3: Have Alice send the message 'Hey Bob, want to play Chess?'.",
            "Expected Result: In Bob's specific chat window, the message should instantly spawn on the screen without clicking refresh. You should also see typing indicators ('User is typing...') when either side is entering a message."
        ]),
        ("Test Module 5: Admin Panel & Moderation", [
            "Objective: Verify the administrative dashboard correctly tracks metrics and deletes user data.",
            "Action 1: Manually change the URL to http://127.0.0.1:5000/admin.",
            "Expected Result: You see a statistical overview: Total Users, Total Hobbies, and Total Reports.",
            "Action 2: Locate Bob in the user table on the admin page and click the 'Delete' button.",
            "Expected Result: Bob is deleted. Any messages Bob sent to Alice are cascaded or safely removed from the system, preventing data orphans."
        ])
    ]
    
    for idx, (test_title, test_steps_list) in enumerate(test_cases, 1):
        doc.add_heading(test_title, level=2)
        for s_idx, step in enumerate(test_steps_list):
            p = doc.add_paragraph(f"{step}")
            if "Expected Result" in step:
                p.style = "Intense Quote"
            elif "Action" in step:
                p.style = "List Bullet"
        doc.add_paragraph("\n")
        
    doc.add_heading("4. Notes on Screenshots", level=1)
    doc.add_paragraph("*(Placeholder Note)*: When taking manual screenshots for submission, simply insert them under the \'Expected Result\' descriptions above for each Module test using the Microsoft Word \'Insert > Pictures\' functionality.")

    output_name = "Deploy_and_Test_Guide_Elaborated.docx"
    doc.save(output_name)
    print(f"Successfully created {output_name}")

if __name__ == "__main__":
    main()
