import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def main():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading("Campus Hobby Connector", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_heading("Deployment and Testing Guide", level=1)
    
    # Deployment Section
    doc.add_heading("1. Deployment Steps", level=2)
    doc.add_paragraph("Follow these instructions to deploy and run the application locally on your machine.")
    
    deploy_steps = [
        ("Prerequisites", "Ensure that Python 3.x is installed on your system."),
        ("Extract and Navigate", "Extract the project folder and open a terminal/command prompt at the root of the project directory."),
        ("Virtual Environment (Optional)", "Run `python -m venv .venv` to create a virtual environment. Activate it using `.\\.venv\\Scripts\\activate` on Windows or `source .venv/bin/activate` on Mac/Linux."),
        ("Install Dependencies", "Run `pip install -r requirements.txt` to install Flask, Flask-SQLAlchemy, Flask-SocketIO, and all other required packages."),
        ("Start the Application", "Execute the command `python app.py`. The database (database.db) will automatically initialize and seed with default hobbies."),
        ("Access the Site", "Open your web browser and navigate to http://127.0.0.1:5000/")
    ]
    
    for step_title, step_desc in deploy_steps:
        p = doc.add_paragraph()
        p.add_run(f"{step_title}: ").bold = True
        p.add_run(step_desc)
        p.style = "List Bullet"
        
    doc.add_paragraph("\n")
    
    # Testing Section
    doc.add_heading("2. Testing Guide", level=2)
    doc.add_paragraph("Below are the core test scenarios to verify the system works correctly.")
    
    test_cases = [
        ("User Authentication", [
            "Navigate to http://127.0.0.1:5000/.",
            "Click 'Register' and create two separate test accounts (e.g., using a standard browser and an Incognito/Private window).",
            "Log out and log back in to verify the login system correctly validates passwords."
        ]),
        ("Profile & Hobbies Setup", [
            "Log in and navigate to your Profile to add or update your bio.",
            "Navigate to the Hobbies section.",
            "Select default hobbies from the checklist (e.g., Photography, Coding) and add custom hobbies via the text field.",
            "Save and check the Dashboard to confirm your hobbies are reflected."
        ]),
        ("Matchmaking Engine", [
            "With both test accounts configured to have at least one intersecting hobby, log into one of the accounts.",
            "Navigate to the 'Matches' tab from the navigation bar.",
            "Verify that the system assigns a proper match percentage based on shared hobbies."
        ]),
        ("Real-time Chat (SocketIO)", [
            "With Test Account A in one window and Test Account B in another, go to the 'Chat' section.",
            "Initiate a chat session between them.",
            "Type a message from A to B. Verify that the 'User is typing...' indicator shows up, and the message appears instantly for User B without reloading the page."
        ]),
        ("Admin Panel", [
            "Navigate directly to http://127.0.0.1:5000/admin.",
            "Verify that you can see system statistics (Total Users, Total Hobbies, Total Reports).",
            "Try deleting a test user and verify that their related database entries (hobbies, messages) are safely removed."
        ])
    ]
    
    for test_title, test_steps_list in test_cases:
        p = doc.add_paragraph()
        p.add_run(test_title).bold = True
        for step in test_steps_list:
            step_p = doc.add_paragraph(f"{step}")
            step_p.style = "List Bullet 2"
            
    output_name = "Deploy_and_Test_Guide.docx"
    doc.save(output_name)
    print(f"Successfully created {output_name}")

if __name__ == "__main__":
    main()
